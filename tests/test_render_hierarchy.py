import tempfile
import unittest
from pathlib import Path

from docx import Document

import render_hierarchy as hierarchy


def make_doc(path: Path, headers, rows):
    doc = Document()
    doc.add_heading("NODES", 1)
    table = doc.add_table(rows=1, cols=len(headers))
    for i, value in enumerate(headers):
        table.rows[0].cells[i].text = value
    for values in rows:
        cells = table.add_row().cells
        for i, value in enumerate(values):
            cells[i].text = value
    doc.save(path)


class HierarchyTests(unittest.TestCase):
    def parse(self, headers, rows):
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "input.docx"
            make_doc(path, headers, rows)
            return hierarchy.read_docx(path)

    def test_aliases_and_valid_tree(self):
        result = self.parse(["Key", "Parent", "Name", "Students"], [
            ["root", "", "Institution", "100"], ["child", "root", "Programme", "25"],
        ])
        self.assertEqual("root", result.root_id)
        self.assertEqual(25, result.nodes[1]["value"])
        self.assertFalse([i for i in result.issues if i.level == "error"])

    def test_duplicate_and_missing_parent_are_removed(self):
        result = self.parse(["ID", "Parent ID", "Label"], [
            ["root", "", "Root"], ["root", "", "Duplicate"], ["orphan", "missing", "Orphan"],
        ])
        self.assertEqual(["root"], [n["id"] for n in result.nodes])
        self.assertEqual(2, len([i for i in result.issues if i.level == "error"]))

    def test_cycle_is_reported(self):
        result = self.parse(["ID", "Parent", "Label"], [
            ["a", "b", "A"], ["b", "a", "B"],
        ])
        self.assertFalse(result.nodes)
        self.assertTrue(any("Circular" in i.message for i in result.issues))

    def test_multiple_roots_get_synthetic_root(self):
        result = self.parse(["ID", "Parent", "Label"], [
            ["a", "", "A"], ["b", "", "B"],
        ])
        self.assertEqual("__wordviz_root__", result.root_id)
        self.assertEqual("__wordviz_root__", result.nodes[1]["parent_id"])
        self.assertTrue(any("roots found" in i.message for i in result.issues))

    def test_unsafe_link_and_invalid_value_are_cleaned(self):
        result = self.parse(["ID", "Label", "Value", "Link"], [
            ["a", "A", "-2", "javascript:alert(1)"],
        ])
        self.assertEqual(1, result.nodes[0]["value"])
        self.assertEqual("", result.nodes[0]["link"])


if __name__ == "__main__":
    unittest.main()

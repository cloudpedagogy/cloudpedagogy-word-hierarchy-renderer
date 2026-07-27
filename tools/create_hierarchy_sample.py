#!/usr/bin/env python3
"""Create the populated hierarchy example Word document."""

from pathlib import Path
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

OUT = Path(__file__).parents[1] / "examples/hierarchy_example.docx"


def shade(cell, fill):
    props = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    props.append(shd)


def set_cell_margins(cell):
    props = cell._tc.get_or_add_tcPr()
    margins = props.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        props.append(margins)
    for side, value in (("top", 80), ("bottom", 80), ("start", 120), ("end", 120)):
        node = OxmlElement(f"w:{side}")
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")
        margins.append(node)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = False
    for i, (header, width) in enumerate(zip(headers, widths)):
        cell = table.rows[0].cells[i]
        cell.width = Inches(width)
        cell.text = header
        shade(cell, "E8EEF5")
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(8.5)
    table.rows[0]._tr.get_or_add_trPr().append(OxmlElement("w:tblHeader"))
    for values in rows:
        cells = table.add_row().cells
        for i, (value, width) in enumerate(zip(values, widths)):
            cells[i].width = Inches(width)
            cells[i].text = str(value)
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cells[i])
            for p in cells[i].paragraphs:
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.0
                for run in p.runs:
                    run.font.size = Pt(8.5)
        trPr = table.rows[-1]._tr.get_or_add_trPr()
        cant = OxmlElement("w:cantSplit")
        trPr.append(cant)
    return table


def build():
    doc = Document()
    sec = doc.sections[0]
    sec.orientation = WD_ORIENT.LANDSCAPE
    sec.page_width, sec.page_height = Inches(11), Inches(8.5)
    sec.top_margin = sec.bottom_margin = Inches(.65)
    sec.left_margin = sec.right_margin = Inches(.65)
    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(10)
    styles["Normal"].paragraph_format.space_after = Pt(6)
    for name, size in (("Title", 22), ("Heading 1", 15)):
        styles[name].font.name = "Calibri"
        styles[name].font.size = Pt(size)
        styles[name].font.color.rgb = RGBColor(46, 116, 181)
    title = doc.add_paragraph("Word-to-Interactive Hierarchy: Sample Input", style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    doc.add_paragraph("Edit the SETTINGS and NODES tables, then run render_hierarchy.py. Parent IDs define the hierarchy; leave Parent ID blank for the root.")
    doc.add_heading("SETTINGS", level=1)
    add_table(doc, ["Setting", "Value"], [
        ["title", "Moodle Course Portfolio"],
        ["subtitle", "Example institutional hierarchy generated from editable Word tables"],
        ["default_layout", "tree"],
        ["allow_layout_switching", "true"],
        ["colour_by", "type"],
        ["size_by", "value"],
        ["show_labels", "true"],
        ["theme", "light"],
    ], [2.2, 7.5])
    doc.add_paragraph()
    doc.add_heading("NODES", level=1)
    rows = [
        ["lshtm", "", "LSHTM", "Institution", "100", "Active", "London School of Hygiene & Tropical Medicine.", "#2457A7", ""],
        ["php", "lshtm", "Public Health and Policy", "Faculty", "55", "Active", "Faculty portfolio.", "", ""],
        ["eph", "lshtm", "Epidemiology and Population Health", "Faculty", "45", "Active", "Faculty portfolio.", "", ""],
        ["mscph", "php", "MSc Public Health", "Programme", "32", "Redesign", "Programme with core and elective modules.", "", ""],
        ["ghp", "php", "MSc Global Health Policy", "Programme", "23", "Active", "Global health policy programme.", "", ""],
        ["epi", "eph", "MSc Epidemiology", "Programme", "28", "Active", "Epidemiology programme.", "", ""],
        ["core", "mscph", "Core modules", "Module group", "18", "Active", "Required modules shared by all students.", "", ""],
        ["elective", "mscph", "Elective modules", "Module group", "14", "Active", "Optional specialist modules.", "", ""],
        ["phm102", "core", "Basic Statistics", "Module", "250", "Active", "Introductory statistical concepts and methods.", "#469D77", "https://www.lshtm.ac.uk/"],
        ["phm201", "core", "Epidemiology", "Module", "230", "Active", "Core epidemiology teaching.", "", ""],
        ["ghm201", "ghp", "Global Health Policy", "Module", "120", "Review", "Policy analysis and global health governance.", "", ""],
        ["epm101", "epi", "Study Design", "Module", "160", "Active", "Study design and critical appraisal.", "", ""],
        ["epm202", "epi", "Advanced Methods", "Module", "90", "Active", "Advanced epidemiological methods.", "", ""],
    ]
    add_table(doc, ["ID", "Parent ID", "Label", "Type", "Value", "Status", "Description", "Colour", "Link"], rows,
              [.72, .82, 1.42, .92, .52, .68, 2.18, .72, .72])
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
